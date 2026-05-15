# -*- coding: utf-8 -*-
"""Extract command/free-speech candidates from doc/requirements.

This script is intentionally conservative: it builds a reviewable corpus and
task matrix, but it does not treat generated negative samples as final oracle.
The output is written under cucumber-agent-testing/debug so it does not pollute
the project root.
"""

from __future__ import annotations

import argparse
import csv
import datetime as _dt
import json
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional


ROOT = Path(__file__).resolve().parents[3]
DEFAULT_INPUT = ROOT / "doc" / "requirements"
DEFAULT_OUTPUT_ROOT = ROOT / "satellite" / "cucumber-agent-testing" / "debug" / "requirements_corpus"


PHRASE_KEYS = [
    "命令词",
    "命令",
    "指令",
    "词条",
    "语料",
    "说法",
    "文本",
    "utterance",
    "sentence",
    "query",
    "phrase",
]
INTENT_KEYS = ["意图", "intent", "功能", "技能", "skill"]
SLOT_KEYS = ["slot", "槽", "参数", "变量"]
ACTION_KEYS = ["动作", "执行", "控制", "action"]
RESPONSE_KEYS = ["播报", "回复", "应答", "tts", "response"]
TYPE_KEYS = ["类型", "分类", "类别", "分组", "category", "tag"]
MAX_PHRASE_CHARS = 48


def read_text_file(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def norm(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def match_key(headers: List[str], keys: List[str]) -> Optional[int]:
    lowered = [h.lower() for h in headers]
    for key in keys:
        key_l = key.lower()
        for idx, header in enumerate(lowered):
            if key_l in header:
                return idx
    return None


def classify_record(source: str, sheet: str, raw: str, row_type: str = "") -> str:
    hay = f"{source} {sheet} {row_type} {raw}".lower()
    if any(k in hay for k in ["自由说", "自由语料", "slot", "意图", "free"]):
        return "free_speech"
    if any(k in hay for k in ["在线", "云端", "asr", "online"]):
        return "online"
    if any(k in hay for k in ["命令词", "指令", "命令", "command"]):
        return "command"
    return "unknown"


def looks_like_non_utterance(text: str) -> bool:
    text = norm(text)
    if not text:
        return True
    if len(text) > MAX_PHRASE_CHARS:
        return True
    bad_tokens = [
        "日志显示",
        "查看版本",
        "喇叭先播报",
        "播报音播放过程中",
        "测试",
        "步骤",
        "预期",
        "设备",
        "安装",
        "下载",
        "联网状态",
    ]
    if any(token in text for token in bad_tokens):
        return True
    if re.match(r"^\d+[、.．]", text):
        return True
    if len(re.findall(r"[，,；;。]", text)) >= 2:
        return True
    return False


def build_record(
    *,
    source_file: Path,
    sheet: str,
    row: int,
    phrase: str,
    intent: str = "",
    slots: str = "",
    action: str = "",
    response: str = "",
    row_type: str = "",
    raw: str = "",
    forced_kind: str = "",
) -> Optional[Dict[str, Any]]:
    phrase = norm(phrase)
    if looks_like_non_utterance(phrase):
        return None
    kind = forced_kind or classify_record(source_file.name, sheet, raw or phrase, row_type)
    return {
        "kind": kind,
        "phrase": phrase,
        "expected_text": phrase,
        "intent": norm(intent),
        "slots": norm(slots),
        "action": norm(action),
        "response": norm(response),
        "type": norm(row_type),
        "source_file": str(source_file.relative_to(ROOT)) if source_file.is_relative_to(ROOT) else str(source_file),
        "sheet": sheet,
        "row": row,
        "raw": norm(raw),
        "review_status": "candidate_needs_review" if kind == "unknown" else "candidate",
    }


def values_matrix(ws: Any) -> List[List[str]]:
    return [[norm(c) for c in row] for row in ws.iter_rows(values_only=True)]


def extract_command_table(path: Path, sheet: str, rows: List[List[str]]) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    current_intent = ""
    for row_no, values in enumerate(rows[1:], start=2):
        if not any(values):
            continue
        if values[0]:
            current_intent = values[0]
        phrase = values[1] if len(values) > 1 else ""
        semantic = values[2] if len(values) > 2 else ""
        record = build_record(
            source_file=path,
            sheet=sheet,
            row=row_no,
            phrase=phrase,
            intent=current_intent,
            action=current_intent,
            response="",
            row_type="command_table",
            raw=" | ".join(values),
            forced_kind="command",
        )
        if record:
            record["semantic"] = semantic
            records.append(record)
    return records


def extract_high_frequency_table(path: Path, sheet: str, rows: List[List[str]]) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    for row_no, values in enumerate(rows[1:], start=2):
        if not any(values):
            continue
        intent = values[0]
        if not intent:
            continue
        for col_no, phrase in enumerate(values[1:], start=2):
            record = build_record(
                source_file=path,
                sheet=sheet,
                row=row_no,
                phrase=phrase,
                intent=intent,
                row_type="high_frequency",
                raw=f"col={col_no} | " + " | ".join(values),
                forced_kind="free_speech",
            )
            if record:
                records.append(record)
    return records


def extract_test_case_table(path: Path, sheet: str, rows: List[List[str]]) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    if not rows:
        return records
    headers = rows[0]
    def idx(name: str) -> int:
        return headers.index(name) if name in headers else -1
    step_i = idx("操作步骤")
    expected_i = idx("预期结果")
    level3_i = idx("用例三级分类")
    level5_i = idx("用例类型")
    name_i = idx("用例名称")
    if step_i < 0:
        return records
    for row_no, values in enumerate(rows[1:], start=2):
        if step_i >= len(values):
            continue
        steps = values[step_i]
        expected = values[expected_i] if expected_i >= 0 and expected_i < len(values) else ""
        level3 = values[level3_i] if level3_i >= 0 and level3_i < len(values) else ""
        case_type = values[level5_i] if level5_i >= 0 and level5_i < len(values) else ""
        case_name = values[name_i] if name_i >= 0 and name_i < len(values) else ""
        raw = " | ".join(values)
        for marker, phrase in re.findall(r"(Wakeup|Asr)#talk#([^#]+)#", steps, flags=re.I):
            marker_l = marker.lower()
            if marker_l == "wakeup":
                forced_kind = "wake"
                intent = "wake"
            else:
                forced_kind = "online" if "在线" in level3 or "在线" in case_name else "command"
                intent = case_type or case_name
            record = build_record(
                source_file=path,
                sheet=sheet,
                row=row_no,
                phrase=phrase,
                intent=intent,
                action=case_type,
                response=expected,
                row_type="test_case_marker",
                raw=raw,
                forced_kind=forced_kind,
            )
            if record:
                records.append(record)
    return records


def extract_xlsx(path: Path) -> List[Dict[str, Any]]:
    try:
        import openpyxl  # type: ignore
    except Exception as exc:  # pragma: no cover - environment dependent
        return [{
            "kind": "error",
            "source_file": str(path),
            "error": f"openpyxl unavailable: {exc}",
        }]

    records: List[Dict[str, Any]] = []
    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    for ws in workbook.worksheets:
        row_values = values_matrix(ws)
        if "命令词表" in ws.title:
            records.extend(extract_command_table(path, ws.title, row_values))
            continue
        if "高频词" in ws.title:
            records.extend(extract_high_frequency_table(path, ws.title, row_values))
            continue
        if row_values and "操作步骤" in row_values[0]:
            records.extend(extract_test_case_table(path, ws.title, row_values))
            continue

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header_idx = None
        headers: List[str] = []
        for idx, row in enumerate(rows[:10]):
            cells = [norm(c) for c in row]
            if any(any(key.lower() in cell.lower() for key in PHRASE_KEYS) for cell in cells):
                header_idx = idx
                headers = cells
                break
        if header_idx is None:
            header_idx = 0
            headers = [f"col_{i + 1}" for i in range(max(len(rows[0]), 1))]

        phrase_i = match_key(headers, PHRASE_KEYS)
        intent_i = match_key(headers, INTENT_KEYS)
        slot_i = match_key(headers, SLOT_KEYS)
        action_i = match_key(headers, ACTION_KEYS)
        response_i = match_key(headers, RESPONSE_KEYS)
        type_i = match_key(headers, TYPE_KEYS)

        for offset, row in enumerate(rows[header_idx + 1 :], start=header_idx + 2):
            values = [norm(c) for c in row]
            if not any(values):
                continue
            if phrase_i is not None and phrase_i < len(values):
                phrase = values[phrase_i]
            else:
                phrase = next((v for v in values if 1 <= len(v) <= 60), "")
            raw = " | ".join(values)
            record = build_record(
                source_file=path,
                sheet=ws.title,
                row=offset,
                phrase=phrase,
                intent=values[intent_i] if intent_i is not None and intent_i < len(values) else "",
                slots=values[slot_i] if slot_i is not None and slot_i < len(values) else "",
                action=values[action_i] if action_i is not None and action_i < len(values) else "",
                response=values[response_i] if response_i is not None and response_i < len(values) else "",
                row_type=values[type_i] if type_i is not None and type_i < len(values) else "",
                raw=raw,
            )
            if record:
                records.append(record)
    return records


def extract_docx(path: Path) -> List[Dict[str, Any]]:
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover - environment dependent
        return [{
            "kind": "error",
            "source_file": str(path),
            "error": f"python-docx unavailable: {exc}",
        }]

    document = docx.Document(str(path))
    lines: List[str] = []
    for paragraph in document.paragraphs:
        text = norm(paragraph.text)
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [norm(cell.text) for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return extract_lines(path, lines)


def extract_delimited(path: Path) -> List[Dict[str, Any]]:
    text = read_text_file(path)
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    try:
        rows = list(csv.reader(text.splitlines(), delimiter=delimiter))
    except Exception:
        return extract_lines(path, text.splitlines())
    if not rows:
        return []
    headers = [norm(c) for c in rows[0]]
    phrase_i = match_key(headers, PHRASE_KEYS)
    intent_i = match_key(headers, INTENT_KEYS)
    slot_i = match_key(headers, SLOT_KEYS)
    action_i = match_key(headers, ACTION_KEYS)
    response_i = match_key(headers, RESPONSE_KEYS)
    type_i = match_key(headers, TYPE_KEYS)
    records = []
    for row_no, row in enumerate(rows[1:], start=2):
        values = [norm(c) for c in row]
        if not any(values):
            continue
        phrase = values[phrase_i] if phrase_i is not None and phrase_i < len(values) else values[0]
        record = build_record(
            source_file=path,
            sheet="csv",
            row=row_no,
            phrase=phrase,
            intent=values[intent_i] if intent_i is not None and intent_i < len(values) else "",
            slots=values[slot_i] if slot_i is not None and slot_i < len(values) else "",
            action=values[action_i] if action_i is not None and action_i < len(values) else "",
            response=values[response_i] if response_i is not None and response_i < len(values) else "",
            row_type=values[type_i] if type_i is not None and type_i < len(values) else "",
            raw=" | ".join(values),
        )
        if record:
            records.append(record)
    return records


def extract_lines(path: Path, lines: Iterable[str]) -> List[Dict[str, Any]]:
    records = []
    for row_no, line in enumerate(lines, start=1):
        text = norm(line)
        if not text or len(text) < 2:
            continue
        if text.startswith(("#", "##", "###")):
            continue
        parts = [norm(p) for p in re.split(r"\t|\s{2,}|[|,，]", text) if norm(p)]
        phrase = ""
        if len(parts) >= 2:
            phrase = parts[0]
        elif re.search(r"[:：]", text):
            phrase = norm(re.split(r"[:：]", text, maxsplit=1)[-1])
        else:
            phrase = text
        if len(phrase) > 60:
            continue
        record = build_record(source_file=path, sheet="text", row=row_no, phrase=phrase, raw=text)
        if record:
            records.append(record)
    return records


def extract_file(path: Path) -> List[Dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return extract_xlsx(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".csv", ".tsv"}:
        return extract_delimited(path)
    if suffix in {".txt", ".md"}:
        return extract_lines(path, read_text_file(path).splitlines())
    return []


def dedupe(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    result = []
    for record in records:
        if record.get("kind") == "error":
            result.append(record)
            continue
        key = (record.get("kind"), record.get("phrase"), record.get("intent"), record.get("slots"))
        if key in seen:
            continue
        seen.add(key)
        result.append(record)
    return result


def generate_variants(records: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    prefixes = ["请", "帮我", "麻烦", "现在"]
    suffixes = ["一下", "吧", "可以吗"]
    variants: List[Dict[str, Any]] = []
    for record in records:
        if record.get("kind") not in {"command", "free_speech", "online"}:
            continue
        phrase = record["phrase"]
        for prefix in prefixes:
            variants.append({
                "base_phrase": phrase,
                "variant": f"{prefix}{phrase}",
                "variant_type": "prefix",
                "expected_intent": record.get("intent", ""),
                "source": record.get("source_file", ""),
                "review_status": "synthetic_candidate_needs_review",
            })
        for suffix in suffixes:
            variants.append({
                "base_phrase": phrase,
                "variant": f"{phrase}{suffix}",
                "variant_type": "suffix",
                "expected_intent": record.get("intent", ""),
                "source": record.get("source_file", ""),
                "review_status": "synthetic_candidate_needs_review",
            })
        for interval_ms in (500, 800, 1000, 1500):
            variants.append({
                "base_phrase": phrase,
                "variant": phrase,
                "variant_type": "oneshot_interval_task",
                "interval_ms": interval_ms,
                "expected_intent": record.get("intent", ""),
                "source": record.get("source_file", ""),
                "review_status": "task_candidate",
            })
    return variants


def write_csv(path: Path, rows: List[Dict[str, Any]]) -> None:
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    headers = sorted({key for row in rows for key in row.keys()})
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        writer.writerows(rows)


def build_report(records: List[Dict[str, Any]], variants: List[Dict[str, Any]], input_dir: Path) -> str:
    counts: Dict[str, int] = {}
    for record in records:
        counts[record.get("kind", "unknown")] = counts.get(record.get("kind", "unknown"), 0) + 1
    lines = [
        "# Requirements Corpus Ingestion Report",
        "",
        f"- input_dir: `{input_dir}`",
        f"- generated_at: `{_dt.datetime.now().isoformat(timespec='seconds')}`",
        f"- records: `{len(records)}`",
        f"- synthetic_variants: `{len(variants)}`",
        "",
        "## Counts",
        "",
    ]
    for key in sorted(counts):
        lines.append(f"- {key}: {counts[key]}")
    lines.extend([
        "",
        "## Next",
        "",
        "- Review `corpus_candidates.csv` and mark final oracle for command/free-speech/online items.",
        "- Review `synthetic_variants.csv`; synthetic negative or口语化 candidates must be confirmed before formal PASS/FAIL.",
        "- Feed reviewed rows into Cucumber action/assertion registry.",
    ])
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-dir", default=str(DEFAULT_INPUT))
    parser.add_argument("--output-dir", default="")
    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    stamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = Path(args.output_dir) if args.output_dir else DEFAULT_OUTPUT_ROOT / stamp
    output_dir.mkdir(parents=True, exist_ok=True)

    records: List[Dict[str, Any]] = []
    if input_dir.exists():
        for path in sorted(p for p in input_dir.rglob("*") if p.is_file() and not p.name.startswith("~$")):
            records.extend(extract_file(path))
    records = dedupe(records)
    variants = generate_variants([r for r in records if r.get("kind") != "error"])

    (output_dir / "corpus_candidates.json").write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "synthetic_variants.json").write_text(json.dumps(variants, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(output_dir / "corpus_candidates.csv", records)
    write_csv(output_dir / "synthetic_variants.csv", variants)
    (output_dir / "ingestion_report.md").write_text(build_report(records, variants, input_dir), encoding="utf-8")
    print(output_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
